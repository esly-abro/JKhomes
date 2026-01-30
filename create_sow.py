from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_document():
    document = Document()

    # Title
    title = document.add_heading('Statement of work: Edu-Explore Website Build', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Metadata
    p = document.add_paragraph()
    p.add_run('Version: ').bold = True
    p.add_run('1.0\n')
    p.add_run('Date: ').bold = True
    p.add_run('January 14, 2026\n')
    p.add_run('Author: ').bold = True
    p.add_run('Navaneetha Krishnan R')

    # 1. Project Information
    document.add_heading('Project information', level=1)
    table = document.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Role'
    hdr_cells[1].text = 'Name'
    hdr_cells[2].text = 'Email and telephone'
    
    # Row 1
    row1 = table.rows[1].cells
    row1[0].text = 'Client point of contact'
    row1[1].text = 'Benita'
    row1[2].text = '[Client Email] [Client Phone]'
    
    # Row 2
    row2 = table.rows[2].cells
    row2[0].text = 'Agency point of contact'
    row2[1].text = 'Navaneetha Krishnan R'
    row2[2].text = '[Agency Email] [Agency Phone]'

    # 2. Project Summary
    document.add_heading('Project summary', level=1)
    summary_table = document.add_table(rows=7, cols=2)
    summary_table.style = 'Table Grid'
    
    data = [
        ('Start date', 'TBD'),
        ('Deadline', 'TBD'),
        ('Project overview', 'The goal of this project is to build a conversion-focused website for Edu-Explore to clearly present services (Career Counselling, Study Abroad, Training) and drive consultation bookings. We aim to revive the inactive domain edu-xplore.com and implement an automated booking flow.'),
        ('Project scope', 'In Scope:\n1. Setup of hosting for edu-xplore.com.\n2. Design and development of Landing Page with conversion CTAs.\n3. Integration of booking engine (Calendly/Google Meet) with Calendar sync.\n4. Creation of service pages: Study Abroad (US, UK, Asia, etc.), Career Counselling, MAX ME Training.\n5. Mobile responsive design.\n\nNot In Scope:\n1. Content writing (text to be provided by client).\n2. Purchase of third-party licenses (hosting, paid booking tools).'),
        ('Project timeline', 'High-level timeline to be confirmed upon booking tool selection and hosting finalization.'),
        ('Project budget', 'TBD'),
        ('Risk, constraints, and assumptions', 'Constraints: Hosting is currently expired and requires immediate renewal.\nAssumptions: Client has access to GoDaddy account for DNS updates.')
    ]

    for i, (key, value) in enumerate(data):
        row = summary_table.rows[i].cells
        row[0].text = key
        row[1].text = value
        row[0].paragraphs[0].runs[0].bold = True

    # 3. Project scope and process
    document.add_heading('Project scope and process', level=1)
    document.add_paragraph('We break the project into 4 key phases to minimize risk through clear checkpoints, reviews, and controls at the beginning and end of a phase.')
    
    phases_list = [
        "Phase 1: Mobilization & Infrastructure",
        "Phase 2: Design & User Flow",
        "Phase 3: Development & Integration",
        "Phase 4: Launch & Handover"
    ]
    for phase in phases_list:
        document.add_paragraph(phase, style='List Bullet')

    # 4. Project Milestones
    document.add_heading('Project milestones', level=1)
    p_italic = document.add_paragraph()
    p_italic.add_run('Please see Project Plan for a detailed project schedule.').italic = True
    
    milestone_table = document.add_table(rows=5, cols=4)
    milestone_table.style = 'Table Grid'
    
    m_headers = ['Milestone', 'Description', 'Start date', 'Due date']
    for i, h in enumerate(m_headers):
        milestone_table.rows[0].cells[i].text = h
        
    m_data = [
        ('Phase 1', 'Hosting setup & Tool selection', 'TBD', 'TBD'),
        ('Phase 2', 'Wireframes & Design Approval', 'TBD', 'TBD'),
        ('Phase 3', 'Functional Website Build', 'TBD', 'TBD'),
        ('Phase 4', 'Go-Live', 'TBD', 'TBD')
    ]
    
    for i, (m, d, s, due) in enumerate(m_data):
        cells = milestone_table.rows[i+1].cells
        cells[0].text = m
        cells[1].text = d
        cells[2].text = s
        cells[3].text = due

    # 5. Project Phases Detail
    document.add_heading('Project phases', level=1)

    # Phase 1
    document.add_heading('Phase 1: Mobilization & Infrastructure', level=2)
    document.add_paragraph('Phase description: Finalizing the technical foundation and tools required for the build.')
    p1_table = document.add_table(rows=3, cols=2)
    p1_table.style = 'Table Grid'
    p1_table.rows[0].cells[0].text = 'Deliverable / task'
    p1_table.rows[0].cells[1].text = 'Description'
    p1_table.rows[1].cells[0].text = 'Hosting Renewal'
    p1_table.rows[1].cells[1].text = 'Assist client in renewing/purchasing hosting for edu-xplore.com.'
    p1_table.rows[2].cells[0].text = 'Tool Selection'
    p1_table.rows[2].cells[1].text = 'Finalize decision between Calendly vs. Google Meet for booking flow.'

    # Phase 2
    document.add_heading('Phase 2: Design & User Flow', level=2)
    document.add_paragraph('Phase description: Defining the visual structure and user journey.')
    p2_table = document.add_table(rows=3, cols=2)
    p2_table.style = 'Table Grid'
    p2_table.rows[0].cells[0].text = 'Deliverable / task'
    p2_table.rows[0].cells[1].text = 'Description'
    p2_table.rows[1].cells[0].text = 'Sitemap & Flow'
    p2_table.rows[1].cells[1].text = 'Define navigation for Study Abroad, Counselling, and Training tabs.'
    p2_table.rows[2].cells[0].text = 'UI Mockups'
    p2_table.rows[2].cells[1].text = 'Create visual mockups for the Landing Page with "Learn More", "Call Now" CTAs.'

    # Phase 3
    document.add_heading('Phase 3: Development & Integration', level=2)
    document.add_paragraph('Phase description: Building the actual website pages and connecting functional elements.')
    p3_table = document.add_table(rows=4, cols=2)
    p3_table.style = 'Table Grid'
    p3_table.rows[0].cells[0].text = 'Deliverable / task'
    p3_table.rows[0].cells[1].text = 'Description'
    p3_table.rows[1].cells[0].text = 'Landing Page Build'
    p3_table.rows[1].cells[1].text = 'Develop primary landing page focused on conversion.'
    p3_table.rows[2].cells[0].text = 'Booking Integration'
    p3_table.rows[2].cells[1].text = 'Implement booking logic, calendar sync, and WhatsApp/Email notification triggers.'
    p3_table.rows[3].cells[0].text = 'Service Pages'
    p3_table.rows[3].cells[1].text = 'Build dedicated pages for MAX ME partnership and Country-wise study programs.'

    # Phase 4
    document.add_heading('Phase 4: Launch & Handover', level=2)
    document.add_paragraph('Phase description: Final testing and deployment to the live server.')
    p4_table = document.add_table(rows=3, cols=2)
    p4_table.style = 'Table Grid'
    p4_table.rows[0].cells[0].text = 'Deliverable / task'
    p4_table.rows[0].cells[1].text = 'Description'
    p4_table.rows[1].cells[0].text = 'Testing'
    p4_table.rows[1].cells[1].text = 'Verify mobile responsiveness and booking flow functionality.'
    p4_table.rows[2].cells[0].text = 'Go-Live'
    p4_table.rows[2].cells[1].text = 'Point domain DNS to new hosting and publish site.'

    # 6. Budget
    document.add_heading('Project budget and billing schedule', level=1)
    p_budget = document.add_paragraph()
    p_budget.add_run('If the project budget has been created, list the payments related to the project below.').italic = True
    
    budget_table = document.add_table(rows=5, cols=4)
    budget_table.style = 'Table Grid'
    b_headers = ['Invoice', 'Milestone/phase', 'Amount', 'Due date']
    for i, h in enumerate(b_headers):
        budget_table.rows[0].cells[i].text = h
    
    b_data = [
        ('1', 'Project Initiation (Upfront)', 'TBD', 'TBD'),
        ('2', 'Design Sign-off', 'TBD', 'TBD'),
        ('3', 'Project Completion', 'TBD', 'TBD'),
        ('Total', '', 'TBD', '')
    ]
    
    for i, (inv, m, amt, due) in enumerate(b_data):
        cells = budget_table.rows[i+1].cells
        cells[0].text = inv
        cells[1].text = m
        cells[2].text = amt
        cells[3].text = due
    
    # Bold the Total row
    budget_table.rows[4].cells[0].paragraphs[0].runs[0].bold = True
    budget_table.rows[4].cells[2].paragraphs[0].runs[0].bold = True

    # 7. General risks
    document.add_heading('General risks, constraints, and assumptions', level=1)
    risks = [
        "Risks: Delays in hosting renewal may push back the start date.",
        "Constraints: The domain edu-xplore.com is valid until 2027, but the site is currently inactive due to expired hosting.",
        "Assumptions: Edu-Explore will provide all necessary logos, images, and text content for the service pages."
    ]
    for r in risks:
        document.add_paragraph(r, style='List Bullet')

    # 8. Change requests
    document.add_heading('Change requests', level=1)
    document.add_paragraph('Any changes to the approved scope (e.g., adding payment gateways, new pages not listed, or complex student portals) will be treated as a Change Request. These will be estimated separately and may incur additional costs and timeline extensions.')

    # 9. Other requirements
    document.add_heading('Other project requirements', level=1)
    reqs = [
        "Access: Developer requires access to the GoDaddy domain management console.",
        "Certification Info: Client to provide specific details regarding the MAX ME Australian certification for the Training page."
    ]
    for req in reqs:
        document.add_paragraph(req, style='List Bullet')

    document.save('Edu_Explore_Statement_of_Work.docx')

if __name__ == "__main__":
    create_document()
